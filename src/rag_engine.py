import os
import json
import docx
from PyPDF2 import PdfReader
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
from dotenv import load_dotenv


from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings




load_dotenv()


# We define the expected LLM output format
class ResumeScore(BaseModel):
    score: float = Field(description="A score between 0 and 100 indicating how well the resume matches the job description.")
    reason: str = Field(description='A brief explation for the given score, highlighting key matches and missing requirements.')


def extract_text_from_pdf(pdf_file) -> str:
    """Extract text from an uploaded PDF file or path."""
    try:
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""
    


def extract_text_from_docx(docx_file) -> str:
    """Extract text from an uploaded DOCX file."""
    try:
        doc = docx.Document(docx_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""
    

def extract_text(uploaded_file) -> str:
    """Extract text from an uploaded file bases on its extension"""
    filename = uploaded_file.name.lower()
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(uploaded_file)
    elif filename.endswith('.docx') or filename.endswith('.doc'):
        return extract_text_from_docx(uploaded_file)
    else:
        print("Unsupported file format.")
        return ""
    

# Ingestion Pipeline
def create_vectorstore(resume_text: str) -> Chroma:
    """Ingestion Pipeline: Chunk resume and create vector store."""
    # Documents: HTML, text files, PDFs (Resume text wrapper)
    doc = Document(page_content=resume_text)

    # Chunking: Split into passages
    text_splitter = RecursiveCharacterTextSplitter(chunk_size = 1000, chunk_overlap=200)
    splits = text_splitter.split_documents([doc])

    # Embedding model & Vector store: Text -> dense vectors & Index chunks + vectors
    embeddings = OpenAIEmbeddings()
    vectorstore = Chroma.from_documents(documents=splits, embedding=embeddings)

    return vectorstore


# Query Pipeline
def retrieve_context(vectorstore: Chroma, query: str) -> str:
    """Query Pipeline: Retrieve top-k similar chunks."""
    # We dynamically get the number of docs in the store (in case it's tiny) to set k
    num_docs = vectorstore._collection.count() if hasattr(vectorstore, "_collection") else 3


    # Query embedding & Retrieval: Top-k similar chunks
    retriever = vectorstore.as_retriever(search_kwargs = {"k": min(3, max(1, num_docs))})
    retrieved_docs = retriever.invoke(query)

    # Prompt assembly: Query + retrieved context
    retrieved_context = "\n\n".join(d.page_content for d in retrieved_docs)
    return retrieve_context


   
def evaluate_with_llm(job_description: str, retrieved_context: str ) -> ResumeScore:
    """LLM Pipeline: Assemble prompt, query LLM, and parse output."""
    # Setup the LLM and Parser
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
    parser = PydanticOutputParser(pydantic_object=ResumeScore)

    # Define the template
    template = """
    You are an expert technical recruiter and hiring manager. Your task is to evaluate a candidate's resume againt a job description.
    Read the job description and the retrieved relevant parts from the candidate's resume carefully.
    Score the candidate out of 100 based on how well theor skills, experience, and education align with the job requirements.

    Job Description (Query):
    {job_description}

    Candiadate Resume (Retrieved Context):
    {retrieved_context}

    Provide your evaluation in the following format:
    {format_instructions}
    """

    prompt = PromptTemplate(
        template=template,
        input_variables=['job_description', 'retrieved_context'],
        partial_variables={'format_instructions': parser.get_format_instructions()}

    )

    # Create the chain: LLM -> Generate grounded answer -> Response
    chain = prompt | llm | parser

    try:
        result = chain.invoke({
            "job_description": job_description,
            "retrieved_context": retrieve_context
        })
        return result
    except Exception as e:
        print(f"Error evaluating resume: {e}")
        return ResumeScore(score=0.0, reason="Error during LLM evaluation.")
    




def grade_resume_against_jd(job_description: str, resume_text: str) -> ResumeScore:
    """
    Orchestrates the RAG pipeline to evaluate a candidate's resume against a job description.
    Returns a Pydantic object containing a score (0-100) and a reason.
    """

    # 1. Ingestion pipeline
    vectorstore = create_vectorstore(resume_text)

    # 2. Query pipeline - Retrieval
    retrieve_context = retrieve_context(vectorstore, job_description)

    # 3. Query pipeline - Generation
    result = evaluate_with_llm(job_description, retrieve_context)

    return result



